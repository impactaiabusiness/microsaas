from flask import Flask, render_template, request, send_file, redirect, url_for, session
from flask_sqlalchemy import SQLAlchemy
from docx import Document
import os
from docx.shared import Pt

app = Flask(__name__)
app.secret_key = 'chave_secreta_segura'

# Configuração do banco de dados SQLite
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///procuracao.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Inicializa o banco de dados
db = SQLAlchemy(app)

# Modelo do banco de dados
class Defesa(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    nome = db.Column(db.String(100), nullable=False)
    nacionalidade = db.Column(db.String(50), nullable=False)
    estado_civil = db.Column(db.String(20), nullable=False)
    profissao = db.Column(db.String(50), nullable=False)
    rg = db.Column(db.String(20), nullable=False)
    orgao = db.Column(db.String(20), nullable=False)
    cpf = db.Column(db.String(14), nullable=False)
    rua = db.Column(db.String(100), nullable=False)
    numero_e_complemento = db.Column(db.String(50), nullable=False)
    bairro = db.Column(db.String(100), nullable=False)
    cidade = db.Column(db.String(100), nullable=False)
    estado = db.Column(db.String(2), nullable=False)
    cep = db.Column(db.String(9), nullable=False)
    ddd = db.Column(db.String(4), nullable=False)
    telefone = db.Column(db.String(15), nullable=False)
    email = db.Column(db.String(150), nullable=False)

# Criar banco de dados
with app.app_context():
    db.create_all()

# Rota para exibir a página inicial
@app.route('/')
def home():
    return render_template('home.html')

# Rota para exibir a página de login
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        usuario = request.form.get('usuario')
        senha = request.form.get('senha')
        if (usuario == 'igor' or usuario == 'vitor' or usuario == 'david') and (senha == 'igor' or senha == 'vitor' or senha == 'david'):
            session['usuario'] = usuario
            return redirect(url_for('procuracao'))  # Alterado para '/procuracao'
        else: 
            return render_template('login.html', erro='Usuário ou senha incorretos!')
    return render_template('login.html')

# Rota para exibir o formulário de procuração
@app.route('/procuracao')  # Alterado para '/procuracao'
def procuracao():
    if 'usuario' not in session:
        return redirect(url_for('login'))
    return render_template('index.html')

# Função para preencher o documento Word corretamente
def preencher_docx(dados):
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    modelo_path = os.path.join(BASE_DIR, "templates", "arquivo1.docx")

    if not os.path.exists(modelo_path):
        raise FileNotFoundError(f"Modelo não encontrado: {modelo_path}")

    doc = Document(modelo_path)

    # Converter todos os valores para maiúsculas
    dados = {chave: valor.strip().upper() for chave, valor in dados.items()}

    # Substituir placeholders nos parágrafos
    for paragrafo in doc.paragraphs:
        texto_modificado = paragrafo.text
        for chave, valor in dados.items():
            placeholder = f"{{{{{chave}}}}}"
            texto_modificado = texto_modificado.replace(placeholder, valor)
        
        if texto_modificado != paragrafo.text:
            paragrafo.clear()
            run = paragrafo.add_run(texto_modificado)
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    # Substituir placeholders dentro das tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                texto_modificado = celula.text
                for chave, valor in dados.items():
                    placeholder = f"{{{{{chave}}}}}"
                    texto_modificado = texto_modificado.replace(placeholder, valor)
                if texto_modificado != celula.text:
                    celula.text = ""
                    run = celula.paragraphs[0].add_run(texto_modificado)
                    run.bold = True
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)

    output_dir = os.path.join(BASE_DIR, "output")
    os.makedirs(output_dir, exist_ok=True)
    nome_arquivo = f"PROCURACAO_{dados.get('nome', 'SEM_NOME')}.docx".replace(" ", "_")
    doc_path = os.path.join(output_dir, nome_arquivo)
    doc.save(doc_path)
    return doc_path

# Rota para processar o formulário e gerar o DOCX
@app.route('/gerar_defesa', methods=['POST'])
def gerar_defesa():
    if 'usuario' not in session:
        return redirect(url_for('login'))
    dados = {key: request.form.get(key) for key in request.form.keys()}
    try:
        docx_path = preencher_docx(dados)
        return send_file(docx_path, as_attachment=True)
    except Exception as e:
        return str(e), 500

# Rota para logout
@app.route('/logout')
def logout():
    session.pop('usuario', None)
    return redirect(url_for('home'))

if __name__ == '__main__':
    app.run(host="0.0.0.0", port=500, debug=True)
